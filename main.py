from pdf2docx import Converter
from docx import Document
import tables.styleTables as st
import tables.parseTables as pt
import textrank.parseDoc as pd
import textrank.textRank as tr
import torch.generate as tg

def run() -> None:
    file_name = get_file_name()
    
    # create xlsx file
    keys = pt.get_keys(file_name)
    conts = pt.get_contents(file_name)

    st.create_xlsx(file_name, keys, conts)

    # get textrank text
    full_text = pd.get_full_text(file_name)
    textrank_text = tr.get_top_scentences(file_name, full_text)
    
    # get t5 text
    doc_headers = pd.get_headers(file_name)
    doc_text = pd.get_text(file_name, doc_headers)


    t5_text = tg.generate_summary(doc_text)
    
    # create doc
    doc = Document()

    doc.add_heading('Syllabus Summaries from Pytorch and TextRank', 0)

    doc.add_heading('T5-Model Summary', level=1)
    doc.add_paragraph(t5_text)

    doc.add_heading('TextRank Summary', level=2)
    doc.add_paragraph(textrank_text)

    document.save('SummarizedSyllabus.docx') 



def convert(pdf_name: str, docx_name: str):
    cv = Converter(pdf_name)
    new = cv.convert(docx_name)      
    cv.close()


def get_file_name() ->str:
    user_file_type = input("Is you syllabus a docx file? (y/n): ")

    if user_file_type == 'y' or user_file_type == 'Y':
        docx_name = input('Please input the docx name (without .docx) and make sure it is in the same directory: ')
        docx = f'{docx_name}.docx'
    else:
        pdf_name = input('Please input the pdf file name and make sure it is in the same directory: ')
        convert_file = f'{pdf_name}.pdf'
        docx = f'{pdf_name}.docx' 
        pdf.convert(f'{pdf_name}.pdf', docx)

    return docx

run()
