from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.opc.constants import RELATIONSHIP_TYPE as rt

links = ['@fullerton.edu', 'zoom.us']
headers = [
    'catalog description',
    'course description',
    'goal',
    'course purpose',
    'required textbook',
    'grading standards and criteria',
    'description of assessed work',
    'late policy',
    'important dates'
]

def get_links(file_name: str) -> dict:
    dict = {} # {rel_id: link} 
    doc = Document(file_name)
    rels = doc.part.rels

    for link in links:
        for rel in rels:
            link_url = rels[rel]._target 
            if rels[rel].reltype == rt.HYPERLINK and link in link_url and 'dsservices' not in link_url and 'StudentITHelpDesk' not in link_url:
                if link == '@fullerton.edu':
                    link_url = link_url[7:]
                dict[rel] = link_url
    return dict


def loop_til_bold(start: int, file_name: str) -> str:
    doc = Document(file_name)
    text = ''
    for i in range(start + 1, len(doc.paragraphs) - 1):
        paragraph = doc.paragraphs[i]
        for run in paragraph.runs:
            if run.bold:
                return text
        text += paragraph.text

def get_runs(file_name: str) -> dict:
    dict = {} # {'header': text}
    doc = Document(file_name)

    for header in headers:
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if header in run.text.lower() and run.bold:
                    txt = loop_til_bold(i, file_name)
                    dict[header] = txt
    print(dict)

get_runs('240.docx')
