from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as rt
import string

links = ['@fullerton.edu', 'zoom.us']
headers = [
    'Catalog Description',
    'Course Description',
    'Course Goal',
    'Student Learning Goals',
    'Course Purpose',
    'Required Textbook',
    'Grading Standards',
    'Late Policy',
    'Late Submission',
    'Important Dates'
]
def clean_whitespace(sections: list) -> list:
    redact = ['  ', '-', '\n\n', '\n', '\t', '\r', '●', '•']
    for i in range(len(sections)):
        text = sections[i]
        for r in redact:
            text = text.replace(r, '')

        sections[i] = text
    s = " ".join(sections)

    return s

#return full text for TextRank
def get_full_text(file_name: str) -> list:
    list = []
    s = ''
    doc = Document(file_name)

    for i, paragraph in enumerate(doc.paragraphs):
        for j, run in enumerate(paragraph.runs):
            if run.bold and j == 0 and len(s) > 5:
                list.append(s)
                s = run.text
            elif len(run.text) > 15:
                s += run.text 
    return list

def get_links(file_name: str) -> dict:
    dict = {} # {rel_id: link} 
    doc = Document(file_name)
    rels = doc.part.rels

    for link in links:
        for rel in rels:
            link_url = rels[rel]._target 
            if rels[rel].reltype == rt.HYPERLINK and link in link_url and 'dsservices' not in link_url and 'StudentITHelpDesk' not in link_url:
                if link == '@fullerton.edu':
                    link_url = link_url[7:]
                dict[rel] = link_url
    return dict

# check if header == paragraph
def is_cleaned(header: str, paragraph: object) -> bool:
    cleaned = header.replace(' ', '')
    return cleaned.lower() in paragraph.text.lower()

# loop until next bold word
def til_bold(file_name: str, start: int) -> str:
    doc = Document(file_name)
    text = ''
    for i in range(start + 1, len(doc.paragraphs) - 1):
        paragraph = doc.paragraphs[i]
        for run in paragraph.runs:
            if run.bold:
                return text
            text += run.text

def get_headers(file_name: str) -> list[list[str, int]]:
    list = [] # ['header', para num]
    doc = Document(file_name)
    for header in headers:
        for i, paragraph in enumerate(doc.paragraphs):
            if is_cleaned(header, paragraph):
                list.append([f'{header}:', i])
                break
            for j, run in enumerate(paragraph.runs):
                if header.lower() in run.text.lower() and j == 0:
                    list.append([f'{header}:', i])
    return list 

def get_text(file_name: str, h: list[list[str, int]]) -> list:
    list = []
    doc = Document(file_name)

    for i in range(len(h)):
        header = h[i][0]
        p_num = h[i][1]
        text = doc.paragraphs[p_num].text
        if len(text) <= len(header) * 2.5:
            list.append(til_bold(file_name, p_num))
        else:
            list.append(text[len(header):])

    lst = clean_whitespace(list)

    return lst

def get_headers_and_text(h: list[list[str, int]], t: list[str]) -> dict:
    dict = {}
    
    for i in range(len(h)):
        header = h[i][0]
        text = t[i]
        dict[header] = text

    return dict

