from pdf2docx import Converter
from docx import Document
import PyPDF2

def convert(pdf_name, docx_name):
    cv = Converter(pdf_name)
    new = cv.convert(docx_name)      
    cv.close()

def is_bold(cell):
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not run.bold:
                return False
    return True
def get_tables(docx):
    keys = []
    contents = []
    data
    doc = Document(docx)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text == 'Grade' or cell.text == 'Percentage' or cell.text == '':
                    continue
                elif is_bold(cell):
                    keys.append(cell.text)
                else:
    for key in keys:
        for content in contents:
            pair = dict(zip(key,content))
            data.append(pair)
            continue

get_tables('335.docx')
