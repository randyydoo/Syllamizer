from docx import Document
import PyPDF2

redact = ['Grade', 'Grades','Percentage', 'A:', 'A-']

def is_bold(cell: 'class') -> bool:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if not run.bold:
                return False
    return True

def get_keys(docx: str) -> list[list[str]]:
    keys = []
    doc = Document(docx)

    for table in doc.tables:
        firstRow = table.rows[0] 
        temp = []
        bold = True
        for cell in firstRow.cells:
            if cell.text == '':
                continue
            elif is_bold(cell) and cell.text not in redact:
                temp.append(cell.text)
        if len(temp) != 0:
            keys.append(temp)
    print(keys)
    print(len(keys))

def get_contents(docx: str) -> list[list[list[str]]]:
    contents = [] 
    doc = Document(docx)

    for table in doc.tables:
        contentTemp = []
        skip, bold, newTable, extendTable = False, True, False, False

        for i, row in enumerate(table.rows):
            tempRow = []

            for j, cell in enumerate(row.cells):
                if cell.text in redact:
                    skip = True

                if j + i == 0:
                    try:
                        prev = int(contents[-1][-1][0])
                        curr = int(cell.text)
                        if curr - prev == 1:
                            extendTable = True
                    except:

                        try:
                            curr = int(cell.text)
                            if curr <= 20:
                                newTable = True
                        except:
                            continue
                if is_bold(cell):
                    continue
                if len(cell.text) != 0 or newTable is True or extendTable is True:
                    if cell.text == '':
                        tempRow.append('None')
                    else:
                        tempRow.append(cell.text)
            if extendTable is True and len(tempRow) > 0:
                contents[-1].append(tempRow)
            elif len(tempRow) != 0:
                contentTemp.append(tempRow)
        if len(contentTemp) != 0 and skip is False:
            contents.append(contentTemp)

    print(contents[3])
    print(len(contents))





def test(docx: str) -> list[list[str]]:
    keys = []  
    contents = [] 
    doc = Document(docx)    
    for table in doc.tables:
        contentTemp, keyTemp = [], []
        new, extend = False, False
        bold, skip= True, False
        for i, row in enumerate(table.rows):
            temp = []
            for j, cell in enumerate(row.cells):
                if j + i == 0:
                    try:
                        prev = int(contents[-1][-1][0])
                        curr = int(cell.text)
                        if curr - prev == 1:
                            extend = True
                    except:
                        try:
                            curr = int(cell.text)
                            if curr <= 20:
                                new = True
                        except:
                            continue
                if cell.text in redact:
                    skip = True
                if is_bold(cell) and i == 0:
                    keyTemp.append(cell.text)
                elif len(keyTemp) != 0 or new is True or extend is True:
                    if cell.text == '':
                        temp.append('None')
                    else:
                        temp.append(cell.text)
            if extend is True and len(temp) > 0:
                contents[-1].append(temp)
            elif len(temp) != 0:
                contentTemp.append(temp)
        if len(contentTemp) != 0 and not skip:
            contents.append(contentTemp)
        if len(keyTemp) != 0 and not skip:
            keys.append(keyTemp)
    print(contents)
    print(len(contents))

test('240.docx')
